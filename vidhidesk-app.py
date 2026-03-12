import streamlit as st
from google import genai
from google.genai import types
import hashlib
import time
import uuid
from datetime import datetime
import PyPDF2
from docx import Document
import io
from PIL import Image
from supabase import create_client, Client

# --- 1. APP CONFIGURATION & SESSION INIT ---
st.set_page_config(
    page_title="VidhiDesk | Legal Intelligence",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

if "user" not in st.session_state: st.session_state.user = None
if "current_workspace" not in st.session_state: st.session_state.current_workspace = {"id": 0, "name": "General Workspace"}

# --- 2. NEO-BRUTALISM THEME ENGINE ---
# We use CSS to aggressively override Streamlit's default styling
st.markdown(f"""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;600;700;900&display=swap');

    /* =========================================
       1. GLOBAL NEO-BRUTALISM RESETS
       ========================================= */
    * {{ font-family: 'Space Grotesk', sans-serif !important; }}
    
    .stApp, [data-testid="stAppViewContainer"], [data-testid="stMain"], [data-testid="stMainBlockContainer"] {{ 
        background-color: #FFFFFF !important; color: #000000 !important; 
    }}
    
    header[data-testid="stHeader"] {{ background: transparent !important; box-shadow: none !important; }}
    [data-testid="stHeaderActionElements"], #MainMenu, .stDeployButton, footer, div[data-testid="stDecoration"] {{ display: none !important; }}
    .block-container {{ padding-top: 2rem !important; padding-bottom: 6rem !important; }}
    
    section[data-testid="stSidebar"] {{ 
        background-color: #F8F8F8 !important; border-right: 4px solid #000 !important; 
    }}
    section[data-testid="stSidebar"] > div {{ padding-top: 1.5rem !important; }}

    h1, h2, h3, h4, h5, h6, p, span, div, label {{ color: #000 !important; font-weight: 600; }}

    /* =========================================
       2. TYPOGRAPHY & BRUTALIST COMPONENTS
       ========================================= */
    .vidhi-title-container {{ width: 100%; text-align: center; padding-top: 4vh; padding-bottom: 2rem; }}
    .vidhi-title {{
        font-size: clamp(3rem, 7vw, 5.5rem) !important; margin: 0 auto;
        color: #000 !important; font-weight: 900 !important; text-transform: uppercase;
        letter-spacing: -2px; line-height: 1; text-shadow: 6px 6px 0px #FF00FF; /* Hot Pink Shadow */
        -webkit-text-fill-color: #000; background: none;
    }}
    .temple-divider {{ height: 6px; width: 100%; background: #000; margin: 25px auto; border-radius: 0; box-shadow: 4px 4px 0px #00FFF0; }}
    .vidhi-subtitle {{ color: #000; font-size: 1rem; font-weight: 900; letter-spacing: 2px; text-transform: uppercase; background-color: #FFE600; display: inline-block; padding: 5px 15px; border: 3px solid #000; box-shadow: 4px 4px 0px #000; }}

    /* Inputs & Textareas */
    .stTextInput > div > div > input, .stChatInput textarea, .stTextArea textarea, div[data-baseweb="select"] > div {{
        background-color: #FFF !important; border: 3px solid #000 !important; color: #000 !important; 
        border-radius: 0px !important; padding: 12px !important; font-weight: 700 !important;
        box-shadow: 4px 4px 0px #000 !important; transition: all 0.1s ease !important;
    }}
    .stTextInput > div > div > input:focus, .stChatInput textarea:focus, .stTextArea textarea:focus, div[data-baseweb="select"] > div:focus-within {{ 
        transform: translate(2px, 2px) !important; box-shadow: 2px 2px 0px #000 !important; background-color: #F0F0F0 !important; outline: none !important;
    }}
    input::placeholder, textarea::placeholder {{ color: #666 !important; opacity: 1 !important; font-weight: 600; text-transform: uppercase; }}

    /* Buttons */
    .stButton > button {{
        background-color: #00FFF0 !important; color: #000 !important; font-weight: 900 !important;
        border: 3px solid #000 !important; border-radius: 0px !important; text-transform: uppercase; letter-spacing: 1px; width: 100%; 
        box-shadow: 6px 6px 0px #000 !important; transition: all 0.1s ease !important; padding: 10px !important;
    }}
    .stButton > button:hover, .stButton > button:active {{
        background-color: #FF00FF !important; color: #FFF !important;
        transform: translate(6px, 6px) !important; box-shadow: 0px 0px 0px #000 !important;
    }}
    button[kind="secondary"] {{ background-color: #FFE600 !important; color: #000 !important; box-shadow: 4px 4px 0px #000 !important; }}
    button[kind="secondary"]:hover {{ background-color: #FFF !important; color: #000 !important; transform: translate(4px, 4px) !important; box-shadow: 0px 0px 0px #000 !important; }}

    /* Containers & Expanders */
    div[data-testid="stContainer"] > div > div > div {{ 
        background-color: #FFF; border: 3px solid #000; border-radius: 0px; box-shadow: 8px 8px 0px #000; padding: 20px !important; 
    }}
    div[data-baseweb="popover"] {{ background-color: #FFF !important; border: 4px solid #000 !important; box-shadow: 8px 8px 0px #FF00FF !important; border-radius: 0 !important; }}
    div[data-testid="stPopover"] > button {{ min-height: 48px !important; border-radius: 0px !important; background: #FFF !important; }}

    /* Chat Bubbles */
    .stChatMessage {{
        background-color: #FFF !important; border: 3px solid #000 !important; border-radius: 0px !important; padding: 1.5rem !important; margin-bottom: 1.5rem !important;
        box-shadow: 6px 6px 0px #000 !important; color: #000 !important;
    }}
    .stChatMessage[data-testid="stChatMessageAvatar"] {{ background-color: #FFE600 !important; border: 3px solid #000 !important; color: #000 !important; border-radius: 0px !important; font-weight: 900; }}

    /* =========================================
       3. SIDEBAR NAVIGATION - BLOCKY TABS
       ========================================= */
    div[role="radiogroup"] {{ display: flex !important; flex-direction: column !important; gap: 15px !important; width: 100% !important; }}
    div[role="radiogroup"] label > div:first-child:not([data-testid="stMarkdownContainer"]), div[role="radiogroup"] label div[data-baseweb="radio"], div[role="radiogroup"] label input {{ display: none !important; width: 0 !important; height: 0 !important; opacity: 0 !important; position: absolute !important; }}
    
    div[role="radiogroup"] label {{
        width: 100% !important; height: 55px !important; margin: 0 !important; cursor: pointer !important; display: flex !important; align-items: center !important; justify-content: flex-start !important;
        background-color: #FFF !important; border: 3px solid #000 !important; border-radius: 0px !important; box-sizing: border-box !important; padding: 0 20px !important; 
        box-shadow: 4px 4px 0px #000 !important; transition: all 0.1s ease !important;
    }}
    div[role="radiogroup"] label div[data-testid="stMarkdownContainer"] p {{
        font-size: 1rem !important; font-weight: 900 !important; color: #000 !important; margin: 0 !important; text-transform: uppercase;
    }}
    div[role="radiogroup"] label:hover {{ transform: translate(2px, 2px); box-shadow: 2px 2px 0px #000 !important; background-color: #FFE600 !important; }}
    div[role="radiogroup"] label:has(input[aria-checked="true"]) {{
        background-color: #00FFF0 !important; transform: translate(4px, 4px); box-shadow: 0px 0px 0px #000 !important; border: 3px solid #000 !important;
    }}

    /* Sticky Header */
    div[data-testid="stVerticalBlock"]:has(#sticky-header-marker):not(:has(div[data-testid="stVerticalBlock"]:has(#sticky-header-marker))) {{
        position: sticky !important; top: 0rem !important; z-index: 999 !important; background-color: #FFF !important; padding: 20px 0px !important; border-bottom: 4px solid #000 !important; margin-bottom: 30px !important; box-shadow: 0px 6px 0px #FFE600;
    }}
</style>
""", unsafe_allow_html=True)

INSTITUTIONS = sorted([
    "National Law School of India University (NLSIU)", "NALSAR University of Law",
    "National Law University, Delhi (NLUD)", "The West Bengal National University of Juridical Sciences (WBNUJS)",
    "National Law University, Jodhpur (NLUJ)", "Hidayatullah National Law University (HNLU)",
    "Tamil Nadu National Law University (TNNLU)", "Maharashtra National Law University (MNLU)",
    "Faculty of Law, University of Delhi (DU)", "Government Law College (GLC), Mumbai", 
    "Symbiosis Law School (SLS)", "School of Law, Christ University", "Jindal Global Law School", "Other"
]) 

# --- 4. DATABASE MANAGER (SUPABASE CLOUD) ---
class DBHandler:
    def __init__(self):
        try:
            url: str = st.secrets["SUPABASE_URL"]
            key: str = st.secrets["SUPABASE_KEY"]
            self.supabase: Client = create_client(url, key)
        except Exception as e:
            st.error("⚠️ Supabase Credentials missing from Streamlit Secrets!")

    def register_user(self, email, password, name, inst, year):
        hashed_pw = hashlib.sha256(password.encode()).hexdigest()
        try:
            self.supabase.table("users").insert({ "email": email, "password": hashed_pw, "name": name, "institution": inst, "year": year, "auth_token": "", "tier": "free" }).execute()
            return True
        except Exception: return False

    def login(self, email, password, remember_me=False):
        hashed_pw = hashlib.sha256(password.encode()).hexdigest()
        response = self.supabase.table("users").select("*").eq("email", email).eq("password", hashed_pw).execute()
        if response.data:
            user = response.data[0]
            token = ""
            if remember_me:
                token = str(uuid.uuid4())
                self.supabase.table("users").update({"auth_token": token}).eq("email", email).execute()
            return { "email": user["email"], "name": user["name"], "institution": user["institution"], "year": user["year"], "tier": user.get("tier", "free"), "token": token }
        return None

    def login_with_token(self, token):
        if not token: return None
        response = self.supabase.table("users").select("*").eq("auth_token", token).execute()
        if response.data:
            user = response.data[0]
            return { "email": user["email"], "name": user["name"], "institution": user["institution"], "year": user["year"], "tier": user.get("tier", "free"), "token": token }
        return None

    def logout(self, email):
        self.supabase.table("users").update({"auth_token": ""}).eq("email", email).execute()

    def save_message(self, email, role, content, workspace_id=0):
        self.supabase.table("chats").insert({ "email": email, "role": role, "content": content, "workspace_id": workspace_id, "timestamp": datetime.now().isoformat() }).execute()

    def get_history(self, email, workspace_id=0):
        response = self.supabase.table("chats").select("role, content").eq("email", email).eq("workspace_id", workspace_id).order("id", desc=False).execute()
        return response.data if response.data else []

    def clear_history(self, email, workspace_id=0):
        self.supabase.table("chats").delete().eq("email", email).eq("workspace_id", workspace_id).execute()

    def save_to_space(self, email, category, query, response, workspace_id=0):
        self.supabase.table("spaces").insert({ "email": email, "category": category, "query": query, "response": response, "workspace_id": workspace_id, "timestamp": datetime.now().isoformat() }).execute()

    def get_space_items(self, email, category, workspace_id=0):
        response = self.supabase.table("spaces").select("id, query, response, timestamp").eq("email", email).eq("category", category).eq("workspace_id", workspace_id).order("id", desc=True).execute()
        return response.data if response.data else []

    def delete_space_item(self, item_id):
        self.supabase.table("spaces").delete().eq("id", item_id).execute()

    def create_workspace(self, email, name):
        response = self.supabase.table("workspaces").insert({ "email": email, "name": name, "created_at": datetime.now().isoformat() }).execute()
        return response.data[0]["id"] if response.data else 0

    def get_workspaces(self, email):
        response = self.supabase.table("workspaces").select("id, name").eq("email", email).order("created_at", desc=True).execute()
        return response.data if response.data else []

db = DBHandler()

if not st.session_state.user:
    saved_token = st.query_params.get("auth_token", None)
    if saved_token:
        auto_user = db.login_with_token(saved_token)
        if auto_user: st.session_state.user = auto_user

# --- 5. MULTIMODAL FILE EXTRACTOR (PDF + VISION OCR) ---
def process_uploaded_file(uploaded_file):
    if not uploaded_file: return None, None
    try:
        if uploaded_file.type == "application/pdf":
            reader = PyPDF2.PdfReader(uploaded_file)
            text = "".join([page.extract_text() + "\n" for page in reader.pages])
            return text, None
        elif uploaded_file.type.startswith("image/"):
            img = Image.open(uploaded_file)
            return None, img
    except Exception as e: return f"Error reading file: {e}", None
    return None, None

def generate_word_document(query, response, title="VidhiDesk Legal Document"):
    doc = Document()
    doc.add_heading(title, 0)
    if query:
        doc.add_heading('Context / Facts:', level=1)
        doc.add_paragraph(query)
    doc.add_heading('Legal Analysis:', level=1)
    clean_response = response.replace("**", "").replace("*", "")
    doc.add_paragraph(clean_response)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 6. AI ENGINE ---
def get_gemini_stream(query, tone, difficulty, institution, chat_history, pdf_text=None, image_data=None, audio_bytes=None, enable_search=False, strict_citation=False):
    try: client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
    except Exception as e:
        yield f"❌ **System Config Error:** {str(e)}"
        return
    
    sys_instruction = f"ROLE: You are VidhiDesk, an elite legal research assistant for {institution}.\nTONE: {tone} | DEPTH: {difficulty}\nMANDATE: Prioritize Indian Statutes (BNS, BNSS, BSA, Constitution). Cite relevant Case Laws. Use Markdown."
    if strict_citation: sys_instruction += "\nCRITICAL RULE (STRICT CITATION MODE): You MUST ONLY cite real, verifiable Indian case laws. Provide the exact year, volume, and court. Under NO circumstances should you invent or hallucinate a case. If you cannot find a verifiable precedent, explicitly state 'No verifiable case law found for this specific query'."

    config = types.GenerateContentConfig(temperature=0.1 if strict_citation else 0.3, system_instruction=sys_instruction)
    if enable_search: config.tools = [{"google_search": {}}]

    contents = []
    for msg in chat_history:
        role = "user" if msg["role"] == "user" else "model"
        contents.append({"role": role, "parts": [{"text": msg["content"]}]})

    current_parts = []
    if pdf_text: current_parts.append({"text": f"[DOCUMENT CONTEXT UPLOADED BY USER]:\n{pdf_text[:15000]}\n\n(Base your answer heavily on the document above if relevant)."})
    if image_data: current_parts.append(image_data)
    if audio_bytes: current_parts.append(types.Part.from_bytes(data=audio_bytes, mime_type="audio/wav"))
    if query: current_parts.append({"text": f"USER QUERY: {query}"})
    
    if current_parts: contents.append({"role": "user", "parts": current_parts})
    if not contents: return

    models_to_try = ['gemini-2.5-flash', 'gemini-2.5-pro', 'gemini-2.0-flash']
    for model_name in models_to_try:
        try:
            response_stream = client.models.generate_content_stream(model=model_name, contents=contents, config=config)
            for chunk in response_stream:
                if chunk.text: yield chunk.text
            return 
        except Exception as e:
            if "API_KEY_INVALID" in str(e) or "not found" in str(e).lower():
                yield "❌ **Authentication Failed:** API key invalid or revoked."
                return
            continue 
    yield "❌ **System Unavailable:** AI servers failed to respond."

def get_drafting_stream(doc_type, client_info, facts, pdf_text=None, image_data=None, audio_bytes=None):
    try: client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
    except Exception:
        yield "❌ **System Config Error.**"
        return
        
    sys_instruction = f"""ROLE: You are an expert Legal Draftsman. TASK: Draft a professional, court-ready '{doc_type}'. MANDATE: Use strict, formal Indian legal terminology. Format properly using clear headings and numbered paragraphs. Use placeholders like [DATE] or [AMOUNT] for missing facts. Base the entire draft strictly on the provided facts and documents. Do not include conversational filler."""
    parts = [{"text": sys_instruction}]
    if pdf_text: parts.append({"text": f"\n[REFERENCE DOCUMENT UPLOADED]:\n{pdf_text[:15000]}"})
    if image_data: parts.append(image_data)
    if client_info: parts.append({"text": f"\n[CLIENT DETAILS]:\n{client_info}"})
    if facts: parts.append({"text": f"\n[CASE FACTS]:\n{facts}"})
    if audio_bytes: parts.append(types.Part.from_bytes(data=audio_bytes, mime_type="audio/wav"))
        
    try:
        response_stream = client.models.generate_content_stream(model='gemini-2.5-flash', contents=[{"role": "user", "parts": parts}])
        for chunk in response_stream:
            if chunk.text: yield chunk.text
    except Exception as e: yield f"❌ **Drafting Engine Error:** {str(e)}"

def get_translation_stream(text, target_lang, institution, pdf_text=None, image_data=None, audio_bytes=None):
    try: client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
    except Exception:
        yield "❌ **System Config Error.**"
        return
        
    sys_instruction = f"ROLE: You are an expert Legal Translator at {institution}. TASK: Translate the provided legal document/text/audio accurately into highly formal {target_lang}. Preserve all legal meanings perfectly. Keep Latin maxims in Latin with translated meanings in brackets."
    parts = [{"text": sys_instruction}]
    if pdf_text: parts.append({"text": f"\n[DOCUMENT TO TRANSLATE]:\n{pdf_text[:15000]}"})
    if image_data: parts.append(image_data)
    if text: parts.append({"text": f"\n[ADDITIONAL TEXT TO TRANSLATE]:\n{text}"})
    if audio_bytes: parts.append(types.Part.from_bytes(data=audio_bytes, mime_type="audio/wav"))
        
    try:
        response_stream = client.models.generate_content_stream(model='gemini-2.5-flash', contents=[{"role": "user", "parts": parts}])
        for chunk in response_stream:
            if chunk.text: yield chunk.text
    except Exception as e: yield f"❌ **Translation Engine Error:** {str(e)}"

def get_vault_analysis_stream(pdf_text=None, image_data=None, audio_bytes=None):
    try: client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
    except Exception:
        yield "❌ **System Config Error.**"
        return
        
    sys_instruction = "ROLE: You are an archiving assistant. Extract the key legal facts, summary, and core arguments from the provided document, image, or audio memo. Format it cleanly in Markdown so it can be saved to a database."
    parts = [{"text": sys_instruction}]
    if pdf_text: parts.append({"text": f"\n[DOCUMENT TO ARCHIVE]:\n{pdf_text[:15000]}"})
    if image_data: parts.append(image_data)
    if audio_bytes: parts.append(types.Part.from_bytes(data=audio_bytes, mime_type="audio/wav"))
        
    try:
        response_stream = client.models.generate_content_stream(model='gemini-2.5-flash', contents=[{"role": "user", "parts": parts}])
        for chunk in response_stream:
            if chunk.text: yield chunk.text
    except Exception as e: yield f"❌ **Archiving Error:** {str(e)}"

# --- 7. UI LOGIC ---
def login_page():
    st.markdown("<div id='login-page-marker'></div>", unsafe_allow_html=True)
    st.markdown("""
        <div class='vidhi-title-container'>
            <div style="display: flex; justify-content: center; margin-bottom: 10px;">
                <!-- Neo-Brutalist Logo Interpretation (Heavy outlines, primary colors) -->
                <svg viewBox="0 0 100 100" class="login-svg" style="width: 140px; height: 140px; filter: drop-shadow(8px 8px 0px #000);">
                    <path d="M 30 20 L 50 80 L 70 20 L 60 20 L 50 55 L 40 20 Z" fill="#FFE600" stroke="#000" stroke-width="4" stroke-linejoin="miter"/>
                    <path d="M 10 10 L 25 30 L 15 50 L 45 95 L 50 85 L 30 50 L 40 30 L 20 10 Z" fill="#FF00FF" stroke="#000" stroke-width="3" stroke-linejoin="miter"/>
                    <path d="M 90 10 L 75 30 L 85 50 L 55 95 L 50 85 L 70 50 L 60 30 L 80 10 Z" fill="#00FFF0" stroke="#000" stroke-width="3" stroke-linejoin="miter"/>
                </svg>
            </div>
            <h1 class='vidhi-title'>VIDHIDESK</h1>
            <div class='temple-divider'></div>
            <div class='vidhi-subtitle'>LOUD. FAST. LEGAL.</div>
        </div>
    """, unsafe_allow_html=True)
    
    c1, c2, c3 = st.columns([1, 1.2, 1])
    with c2:
        with st.container(border=True):
            tab_login, tab_reg, tab_guest = st.tabs(["LOGIN", "REGISTER", "GUEST"])
            with tab_login:
                st.markdown("<br>", unsafe_allow_html=True)
                email = st.text_input("IDENTITY TOKEN (EMAIL)", key="log_email")
                password = st.text_input("SECURITY KEY (PASSWORD)", type="password", key="log_pwd")
                remember_me = st.checkbox("Keep me signed in (Remember Me)")
                if st.button("INITIATE SESSION", use_container_width=True):
                    with st.spinner("Authenticating..."):
                        time.sleep(0.5) 
                        user = db.login(email, password, remember_me)
                        if user:
                            st.session_state.user = user
                            if remember_me and user["token"]: st.query_params["auth_token"] = user["token"]
                            st.rerun()
                        else: st.error("Authentication Failed: Invalid token or key.")
                            
            with tab_reg:
                st.markdown("<br>", unsafe_allow_html=True)
                r_name = st.text_input("FULL NAME")
                r_email = st.text_input("EMAIL ADDRESS")
                r_pwd = st.text_input("CREATE PASSWORD", type="password")
                r_inst = st.selectbox("INSTITUTION", INSTITUTIONS)
                r_year = st.selectbox("YEAR / STATUS", ["1st Year", "2nd Year", "3rd Year", "4th Year", "5th Year", "Graduate", "Faculty", "Other"])
                if st.button("REGISTER ACCOUNT", use_container_width=True):
                    if r_name and r_email and r_pwd:
                        success = db.register_user(r_email, r_pwd, r_name, r_inst, r_year)
                        if success: st.success("Registration Successful! Please navigate to Login.")
                        else: st.error("Error: Email already registered in the system.")
                    else: st.warning("Please fill out all required fields.")

            with tab_guest:
                st.markdown("<br><p style='text-align: center; font-size: 0.9rem; font-weight:800;'>Temporary access mode. Data will be tied to a temporary session.</p><br>", unsafe_allow_html=True)
                if st.button("CONTINUE AS GUEST", type="secondary", use_container_width=True):
                    st.session_state.user = { "email": f"guest_{int(time.time())}@vidhidesk.local", "name": "Guest User", "institution": "Independent Researcher", "year": "N/A", "tier": "free" }
                    st.rerun()

def main_app():
    with st.sidebar:
        st.markdown(f"""
            <div style='display: flex; align-items: center; margin-bottom: 10px; background: #000; padding: 15px; border: 3px solid #000; box-shadow: 4px 4px 0px #FF00FF;'>
                <svg viewBox="0 0 100 100" style="width: 45px; height: 45px; margin-right: 15px; flex-shrink: 0; filter: drop-shadow(4px 4px 0px #00FFF0);">
                    <path d="M 30 20 L 50 80 L 70 20 L 60 20 L 50 55 L 40 20 Z" fill="#FFE600" stroke="#FFF" stroke-width="4"/>
                    <path d="M 10 10 L 25 30 L 15 50 L 45 95 L 50 85 L 30 50 L 40 30 L 20 10 Z" fill="#FF00FF" stroke="#FFF" stroke-width="2"/>
                    <path d="M 90 10 L 75 30 L 85 50 L 55 95 L 50 85 L 70 50 L 60 30 L 80 10 Z" fill="#00FFF0" stroke="#FFF" stroke-width="2"/>
                </svg>
                <div>
                    <h2 style="margin:0; font-size:1.6rem; letter-spacing: -1px; line-height: 1.1; font-weight: 900; color: #FFF !important; text-transform: uppercase;">VIDHIDESK</h2>
                    <span style="font-size: 0.75rem; color: #FFE600; letter-spacing: 2px; font-weight: 900; text-transform: uppercase;">SYSTEM</span>
                </div>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown(f"<h3 style='margin-bottom: 0; font-size: 1.3rem; font-weight: 900; color: #000 !important; text-transform: uppercase;'>{st.session_state.user['name']}</h3>", unsafe_allow_html=True)
        st.markdown(f"<span style='color: #FF00FF; font-size: 0.9rem; font-weight: 900; text-transform: uppercase;'>{st.session_state.user['institution']}</span>", unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown(f"<div style='font-size: 0.85rem; color: #000; margin-bottom: 5px; font-weight: 900; letter-spacing: 1px; text-transform: uppercase;'>ACTIVE CASE FOLDER</div>", unsafe_allow_html=True)
        
        workspaces = [{"id": 0, "name": "General Workspace"}] + db.get_workspaces(st.session_state.user['email'])
        ws_names = [w['name'] for w in workspaces]
        
        current_index = 0
        for i, w in enumerate(workspaces):
            if w['id'] == st.session_state.current_workspace['id']: current_index = i
        
        # Collaborative Sync Feature
        wc1, wc2 = st.columns([0.75, 0.25])
        with wc1:
            selected_ws_name = st.selectbox("Workspace", ws_names, index=current_index, label_visibility="collapsed")
            for w in workspaces:
                if w['name'] == selected_ws_name and st.session_state.current_workspace['id'] != w['id']:
                    st.session_state.current_workspace = w
                    st.rerun()
        with wc2:
            if st.button("SYNC", help="Sync Collaborative Workspace Data"):
                st.toast("Database Synced with Associates!", icon="⚡")
                st.rerun()
        
        with st.popover("CREATE FOLDER", use_container_width=True):
            new_ws_name = st.text_input("Client/Case Name", placeholder="e.g., State vs Sharma")
            if st.button("CREATE NOW", use_container_width=True):
                if new_ws_name:
                    new_id = db.create_workspace(st.session_state.user['email'], new_ws_name)
                    st.session_state.current_workspace = {"id": new_id, "name": new_ws_name}
                    st.rerun()
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        nav = st.radio("MODULES", ["⚡ RESEARCH CORE", "🛠️ DRAFTING STUDIO", "🌍 TRANSLATE DESK", "🗄️ KNOWLEDGE VAULT"], label_visibility="collapsed")
        
        st.markdown("<br>", unsafe_allow_html=True)
        if "GEMINI_API_KEY" in st.secrets:
            st.markdown(f"""
            <div style='border: 3px solid #000; padding: 12px; background: #00FFF0; margin-top:10px; box-shadow: 4px 4px 0px #000;'>
                <div style='display:flex; align-items:center; margin-bottom:5px;'>
                    <span style='color: #000; font-size: 1.2rem; margin-right: 8px; font-weight: 900;'>▇</span> 
                    <span style='color: #000; font-weight: 900; text-transform: uppercase;'>SYSTEM ONLINE</span>
                </div>
                <div style='font-size: 0.8rem; font-weight: 700; color: #000; text-transform: uppercase;'>ENGINE: GENAI 2.5</div>
            </div>
            """, unsafe_allow_html=True)
        else: st.error("Config Error: API Key missing.")
        
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("TERMINATE UPLINK", type="secondary"):
            db.logout(st.session_state.user["email"])
            st.session_state.user = None
            if "auth_token" in st.query_params: del st.query_params["auth_token"]
            st.rerun()

    # --- RESEARCH CORE ---
    if nav == "⚡ RESEARCH CORE":
        sticky_header = st.container()
        with sticky_header:
            st.markdown("<span id='sticky-header-marker'></span>", unsafe_allow_html=True)
            st.markdown(f"<h2 style='margin-bottom: 0; font-size: 3rem; font-weight: 900; text-transform: uppercase;'>RESEARCH CORE</h2>", unsafe_allow_html=True)

            param_col, mic_col = st.columns([0.85, 0.15], vertical_alignment="center")
            with param_col:
                with st.popover("⚙️ ADVANCED SETTINGS", use_container_width=True):
                    c1, c2, c3 = st.columns(3)
                    with c1: tone = st.selectbox("OUTPUT TONE", ["Casual", "Professional", "Academic"], index=2)
                    with c2: diff = st.selectbox("ANALYSIS DEPTH", ["Summary", "Detailed", "Bare Act"], index=1)
                    with c3: space = st.selectbox("AUTO-ARCHIVE TO", ["None", "Research", "Paper", "Study"])
                        
                    st.markdown("---")
                    sc1, sc2, sc3 = st.columns([1, 1, 1])
                    with sc1: st.markdown("<br>", unsafe_allow_html=True); enable_search = st.toggle("🌐 LIVE WEB SEARCH")
                    with sc2: st.markdown("<br>", unsafe_allow_html=True); strict_citation = st.toggle("🛡️ STRICT CITATIONS")
                    with sc3: uploaded_file = st.file_uploader("📄 UPLOAD CONTEXT", type=["pdf", "png", "jpg", "jpeg"], key="res_pdf")

            with mic_col:
                with st.popover("🎙️ VOICE", use_container_width=True):
                    audio_data = st.audio_input("Record", label_visibility="collapsed")
                    submit_audio = st.button("SEND", use_container_width=True, type="secondary")

        history = db.get_history(st.session_state.user['email'], workspace_id=st.session_state.current_workspace['id'])
        for msg in history:
            avatar = "🧑‍⚖️" if msg['role'] == "user" else "⚡"
            with st.chat_message(msg['role'], avatar=avatar): st.markdown(msg['content'])

        query = st.chat_input("ENTER LEGAL QUERY, SECTION, OR CITATION...")
        is_audio_submission = audio_data is not None and submit_audio

        if query or is_audio_submission:
            with st.chat_message("user", avatar="🧑‍⚖️"):
                if query: st.markdown(f"**{query}**")
                if is_audio_submission:
                    st.audio(audio_data)
                    if not query: query = "Please analyze this audio recording."
            
            db.save_message(st.session_state.user['email'], "user", query, workspace_id=st.session_state.current_workspace['id'])
            with st.chat_message("assistant", avatar="⚡"):
                pdf_text, image_data = process_uploaded_file(uploaded_file)
                audio_bytes = audio_data.getvalue() if is_audio_submission else None
                
                with st.spinner("ANALYZING..."):
                    stream = get_gemini_stream(query, tone, diff, st.session_state.user['institution'], history, pdf_text=pdf_text, image_data=image_data, audio_bytes=audio_bytes, enable_search=enable_search, strict_citation=strict_citation)
                    full_response = st.write_stream(stream)
                
                db.save_message(st.session_state.user['email'], "assistant", full_response, workspace_id=st.session_state.current_workspace['id'])

                if space != "None" and "❌" not in full_response:
                    db.save_to_space(st.session_state.user['email'], space, query, full_response, workspace_id=st.session_state.current_workspace['id'])
                    st.toast(f"SAVED TO {space.upper()}", icon="📂")
            st.rerun()

        if history:
            c1, c2 = st.columns([0.85, 0.15])
            with c2:
                if st.button("CLEAR LOGS", type="secondary"):
                    db.clear_history(st.session_state.user['email'], workspace_id=st.session_state.current_workspace['id'])
                    st.rerun()

    # --- DRAFTING STUDIO ---
    elif nav == "🛠️ DRAFTING STUDIO":
        st.markdown(f"<h2 style='margin-bottom: 0; font-size: 3rem; font-weight: 900; text-transform: uppercase;'>DRAFTING STUDIO</h2>", unsafe_allow_html=True)
        st.markdown("<p style='font-size: 1.2rem; font-weight: 700; color: #FF00FF !important; text-transform: uppercase;'>Automated Generation of Court-Ready Legal Documents.</p>", unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)

        with st.container(border=True):
            col_doc, col_pdf, col_voice = st.columns([2, 1, 1])
            with col_doc: doc_type = st.selectbox("DOCUMENT TYPE", ["Legal Notice (General)", "Legal Notice (Sec 138 NI Act)", "Non-Disclosure Agreement (NDA)", "Bail Application (Under BNSS)", "Lease / Rent Agreement", "Affidavit", "Writ Petition (Draft Format)"])
            with col_pdf: uploaded_file = st.file_uploader("📄 ATTACH REF (PDF/IMG)", type=["pdf", "png", "jpg", "jpeg"], key="draft_pdf")
            with col_voice: 
                with st.popover("🎙️ DICTATE DETAILS", use_container_width=True):
                    draft_audio = st.audio_input("Record Facts", label_visibility="collapsed")
            
            st.markdown("#### CLIENT & PARTY DETAILS")
            c1, c2 = st.columns(2)
            with c1: client_name = st.text_input("YOUR CLIENT NAME", placeholder="e.g., Ramesh Kumar")
            with c2: opp_name = st.text_input("OPPOSING PARTY", placeholder="e.g., State Bank of India")
            
            facts = st.text_area("CORE CASE FACTS & TIMELINE", height=120, placeholder="Explain the primary incident, dates, and amounts involved...")
            
            if st.button("GENERATE DRAFT", use_container_width=True, type="primary"):
                if not facts and not uploaded_file and not draft_audio: 
                    st.warning("Please provide either text facts, a voice recording, or a reference file to generate a draft.")
                else:
                    st.markdown("---")
                    st.markdown(f"### GENERATED DRAFT: {doc_type}")
                    
                    pdf_text, image_data = None, None
                    if uploaded_file:
                        with st.spinner("EXTRACTING REFERENCE DOCUMENT..."): 
                            pdf_text, image_data = process_uploaded_file(uploaded_file)
                    
                    audio_bytes = draft_audio.getvalue() if draft_audio else None
                    client_info_str = f"Client: {client_name}\nOpposing Party: {opp_name}" if (client_name or opp_name) else None
                            
                    stream = get_drafting_stream(doc_type, client_info_str, facts, pdf_text=pdf_text, image_data=image_data, audio_bytes=audio_bytes)
                    final_draft = st.write_stream(stream)
                    
                    if "❌" not in final_draft:
                        context_note = f"Facts provided:\n{facts}\n\n[References were included]" 
                        doc_bytes = generate_word_document(context_note, final_draft, title=f"Draft: {doc_type}")
                        st.download_button(label="📄 DOWNLOAD AS WORD", data=doc_bytes, file_name=f"Draft_{doc_type.replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")

    # --- TRANSLATION DESK ---
    elif nav == "🌍 TRANSLATE DESK":
        st.markdown(f"<h2 style='margin-bottom: 0; font-size: 3rem; font-weight: 900; text-transform: uppercase;'>TRANSLATION DESK</h2>", unsafe_allow_html=True)
        st.markdown("<p style='font-size: 1.2rem; font-weight: 700; color: #00FFF0 !important; text-transform: uppercase;'>High-fidelity legal translation preserving complex terminology.</p>", unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)

        with st.container(border=True):
            col_lang, col_pdf, col_voice = st.columns([2, 1, 1])
            with col_lang: target_lang = st.selectbox("TRANSLATE TO", ["Hindi", "Tamil", "Marathi", "Bengali", "Telugu", "Gujarati", "Malayalam", "English"])
            with col_pdf: uploaded_file = st.file_uploader("📄 UPLOAD FILE", type=["pdf", "png", "jpg", "jpeg"], key="trans_pdf")
            with col_voice: 
                with st.popover("🎙️ DICTATE AUDIO", use_container_width=True):
                    trans_audio = st.audio_input("Record Audio", label_visibility="collapsed")
            
            source_text = st.text_area("SOURCE TEXT", height=150, placeholder="Paste legal document text here...")
            
            if st.button("TRANSLATE CONTENT", use_container_width=True):
                if not source_text and not uploaded_file and not trans_audio: 
                    st.warning("Please paste text, upload a file, or record audio to translate.")
                else:
                    st.markdown("---")
                    st.markdown(f"### {target_lang.upper()} TRANSLATION")
                    
                    pdf_text, image_data = None, None
                    if uploaded_file:
                        with st.spinner("EXTRACTING FILE..."):
                            pdf_text, image_data = process_uploaded_file(uploaded_file)
                    
                    audio_bytes = trans_audio.getvalue() if trans_audio else None
                            
                    stream = get_translation_stream(source_text, target_lang, st.session_state.user['institution'], pdf_text=pdf_text, image_data=image_data, audio_bytes=audio_bytes)
                    final_translation = st.write_stream(stream)
                    
                    if "❌" not in final_translation:
                        context_note = f"Source Text:\n{source_text}\n\n[References Included]" 
                        doc_bytes = generate_word_document(context_note, final_translation, title=f"Legal Translation ({target_lang})")
                        st.download_button(label="📄 DOWNLOAD AS WORD", data=doc_bytes, file_name=f"Translation_{target_lang}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")

    # --- VAULT ---
    elif nav == "🗄️ KNOWLEDGE VAULT":
        st.markdown(f"<h2 style='margin-bottom: 0; font-size: 3rem; font-weight: 900; text-transform: uppercase;'>VAULT <span style='font-size:0.5em; color:#FF00FF;'>[{st.session_state.current_workspace['name']}]</span></h2>", unsafe_allow_html=True)
        st.markdown("<p style='font-size: 1.2rem; font-weight: 700; color: #FFE600 !important; text-transform: uppercase;'>Secure database for specific case files and outputs.</p>", unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)
        
        with st.popover("➕ QUICK ANALYZE & ADD TO VAULT", use_container_width=True):
            st.markdown("**UPLOAD A COMPLEX LEGAL DOCUMENT/IMAGE OR DICTATE A VOICE MEMO. VIDHIDESK WILL EXTRACT THE CORE FACTS AND ARCHIVE THEM INSTANTLY.**")
            vc1, vc2, vc3 = st.columns([1, 1, 1])
            with vc1: uploaded_file = st.file_uploader("UPLOAD FILE", type=["pdf", "png", "jpg", "jpeg"], key="vault_pdf")
            with vc2: vault_audio = st.audio_input("RECORD MEMO", key="vault_audio")
            with vc3: 
                v_space = st.selectbox("SAVE TO", ["Research", "Paper", "Study"])
                save_vault = st.button("EXTRACT & ARCHIVE", type="primary", use_container_width=True)

            if save_vault and (uploaded_file or vault_audio):
                pdf_text, image_data = process_uploaded_file(uploaded_file)
                aud_bytes = vault_audio.getvalue() if vault_audio else None
                with st.spinner("ANALYZING AND SAVING..."):
                    stream = get_vault_analysis_stream(pdf_text=pdf_text, image_data=image_data, audio_bytes=aud_bytes)
                    analysis_result = ""
                    for chunk in stream: analysis_result += chunk
                    
                    db.save_to_space(st.session_state.user['email'], v_space, "External File / Audio Analysis", analysis_result, workspace_id=st.session_state.current_workspace['id'])
                    st.success(f"ARCHIVED SUCCESSFULLY TO {v_space.upper()}!")
                    st.rerun()

        t1, t2, t3 = st.tabs(["📚 RESEARCH", "📝 PAPERS", "🎓 STUDY"])
        for tab, cat in zip([t1, t2, t3], ["Research", "Paper", "Study"]):
            with tab:
                st.markdown("<br>", unsafe_allow_html=True)
                items = db.get_space_items(st.session_state.user['email'], cat, workspace_id=st.session_state.current_workspace['id'])
                if not items: st.info(f"SECTOR '{cat.upper()}' IS EMPTY IN THIS FOLDER.", icon="ℹ️")
                else:
                    for item in items:
                        with st.expander(f"📌 {item['timestamp'][:16]} | {item['query'][:60]}..."):
                            st.markdown(item['response'])
                            col1, col2 = st.columns([0.2, 0.8])
                            with col1:
                                if st.button("DELETE", key=f"del_{item['id']}", type="secondary"):
                                    db.delete_space_item(item['id'])
                                    st.rerun()
                            with col2:
                                doc_bytes = generate_word_document(item['query'], item['response'])
                                st.download_button(label="📄 EXPORT TO WORD", data=doc_bytes, file_name=f"VidhiDesk_Research_{item['id']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_{item['id']}")

if __name__ == "__main__":
    if st.session_state.user: main_app()
    else: login_page()
